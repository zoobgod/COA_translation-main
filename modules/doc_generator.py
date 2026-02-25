"""
Word document generation module.

Creates a formatted Russian-language COA Word document with a **fixed
predefined structure**. Each output document always contains the same
sections in the same order, regardless of the original PDF layout.

Supports two modes:
    1. **User-uploaded template** — a .docx file that contains Jinja2
       placeholders (e.g. ``{{ product_name }}``, ``{{ test_results }}``).
       Rendered via docxtpl.
    2. **Built-in fixed structure** — generated from scratch via python-docx
       using the section definitions in ``coa_structure.py``.
"""

import io
import logging
import os
import re
from difflib import SequenceMatcher
from datetime import datetime

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.text.paragraph import Paragraph
from docxtpl import DocxTemplate

from modules.coa_structure import (
    COA_SECTIONS,
    COA_FIELD_KEYS,
    COA_FIELD_LABELS,
    COA_FIELD_IS_TABLE,
)

logger = logging.getLogger(__name__)

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates")

SECTION_ALIASES = {
    "document_title": [
        "наименование документа",
        "certificate of analysis",
        "coa",
        "document title",
        "название документа",
        "title",
    ],
    "company_info": [
        "информация о компании",
        "manufacturer",
        "supplier",
        "company",
        "адрес",
        "контакт",
    ],
    "product_name": [
        "наименование продукта",
        "product name",
        "product",
        "inn",
        "generic name",
    ],
    "product_details": [
        "сведения о продукте",
        "cas",
        "molecular formula",
        "molecular weight",
        "composition",
        "specification",
        "grade",
    ],
    "batch_info": [
        "информация о серии",
        "batch",
        "lot",
        "mfg",
        "expiry",
        "retest",
        "series",
    ],
    "storage_conditions": [
        "условия хранения",
        "storage",
        "хранить",
        "temperature",
        "protect from",
    ],
    "test_results": [
        "результаты испытаний",
        "test results",
        "analysis",
        "assay",
        "impurities",
        "parameter",
        "method",
    ],
    "conclusion": [
        "заключение",
        "conclusion",
        "complies",
        "release",
        "disposition",
    ],
    "signatures": [
        "подписи",
        "signature",
        "approved by",
        "qa",
        "qc",
        "authorised",
    ],
    "notes": [
        "примечания",
        "notes",
        "remark",
        "comment",
        "дополнительно",
    ],
}


# =========================================================================
# Public API
# =========================================================================

def generate_structured_doc(
    sections: dict,
    original_filename: str,
    extraction_method: str,
    model_used: str,
    user_template_bytes: bytes | None = None,
) -> bytes:
    """
    Generate a Word document from structured (section-keyed) translation data.

    Args:
        sections: dict mapping COA field keys → translated Russian content.
                  Table fields contain list[list[str]].
        original_filename: Name of the source PDF.
        extraction_method: Which extraction method produced the source text.
        model_used: OpenAI model used for the translation.
        user_template_bytes: Optional .docx template uploaded by the user.
                             Must contain Jinja2 placeholders matching the
                             COA field keys.

    Returns:
        bytes of the generated .docx file.
    """
    if user_template_bytes:
        return _render_user_template(
            sections, original_filename, extraction_method, model_used,
            user_template_bytes,
        )

    return _generate_fixed_structure(
        sections, original_filename, extraction_method, model_used,
    )


def generate_doc_from_template(
    translated_text: str,
    original_filename: str,
    extraction_method: str,
    model_used: str,
) -> bytes:
    """
    Legacy entry point — builds a fixed-structure document from plain
    translated text (no section mapping).  All translated content is placed
    in the "Результаты / Содержание" area.
    """
    sections = {k: "" for k in COA_FIELD_KEYS}
    sections["notes"] = translated_text
    return _generate_fixed_structure(
        sections, original_filename, extraction_method, model_used,
    )


# =========================================================================
# User-uploaded template rendering (docxtpl)
# =========================================================================

def _render_user_template(
    sections: dict,
    original_filename: str,
    extraction_method: str,
    model_used: str,
    template_bytes: bytes,
) -> bytes:
    """Render a user-provided .docx template with docxtpl."""
    if not _template_has_jinja_placeholders(template_bytes):
        logger.warning(
            "User template has no Jinja placeholders; using structural fallback"
        )
        return _inject_translation_into_template(
            sections,
            original_filename,
            extraction_method,
            model_used,
            template_bytes,
        )

    doc = DocxTemplate(io.BytesIO(template_bytes))

    # Build the context — include every section key plus metadata
    context = dict(sections)

    # Flatten table data to a multi-line string for simple {{ }} placeholders
    for key in COA_FIELD_KEYS:
        if COA_FIELD_IS_TABLE.get(key) and isinstance(sections.get(key), list):
            rows = sections[key]
            context[key] = _table_to_text(rows)

    context.update({
        "original_filename": original_filename,
        "translation_date": datetime.now().strftime("%d.%m.%Y"),
        "model_used": model_used,
        "extraction_method": extraction_method,
    })

    try:
        doc.render(context)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        rendered = buf.getvalue()
    except Exception as e:
        logger.warning("docxtpl render failed, using fallback template mode: %s", e)
        return _inject_translation_into_template(
            sections,
            original_filename,
            extraction_method,
            model_used,
            template_bytes,
        )

    if not _rendered_template_has_translated_content(rendered, sections):
        logger.warning(
            "Rendered template appears to miss translated content; "
            "using structural fallback"
        )
        return _inject_translation_into_template(
            sections,
            original_filename,
            extraction_method,
            model_used,
            template_bytes,
        )

    return rendered


# =========================================================================
# Built-in fixed-structure generation (python-docx)
# =========================================================================

def _generate_fixed_structure(
    sections: dict,
    original_filename: str,
    extraction_method: str,
    model_used: str,
) -> bytes:
    """Build a professionally formatted COA document from scratch."""
    doc = Document()

    # -- Default style -----------------------------------------------------
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    # -- Page margins ------------------------------------------------------
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    # -- Title block -------------------------------------------------------
    _add_title_block(doc)

    # -- Metadata block ----------------------------------------------------
    _add_metadata_block(doc, original_filename, extraction_method, model_used)

    # -- Horizontal rule ---------------------------------------------------
    _add_horizontal_rule(doc)

    # -- Fixed sections, in order ------------------------------------------
    for key, label, _desc, is_table in COA_SECTIONS:
        value = sections.get(key, "")

        # Skip completely empty sections (except test_results which is core)
        if not value and key != "test_results":
            continue

        # Section heading
        _add_section_heading(doc, label)

        if is_table and isinstance(value, list) and len(value) > 0:
            _add_results_table(doc, value)
        elif is_table and isinstance(value, str) and value.strip():
            # Fallback: table came back as text (pipe-delimited)
            _add_text_paragraph(doc, value)
        elif isinstance(value, str) and value.strip():
            _add_text_paragraph(doc, value)
        else:
            # Empty placeholder
            _add_text_paragraph(doc, "—")

    # -- Footer rule -------------------------------------------------------
    _add_horizontal_rule(doc)

    # -- Disclaimer --------------------------------------------------------
    _add_disclaimer(doc)

    # -- Serialise ---------------------------------------------------------
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# =========================================================================
# Formatting helpers
# =========================================================================

def _add_title_block(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("СЕРТИФИКАТ АНАЛИЗА")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(8)
    r = sub.add_run("(Перевод на русский язык)")
    r.font.size = Pt(11)
    r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(100, 100, 100)


def _add_metadata_block(
    doc: Document,
    original_filename: str,
    extraction_method: str,
    model_used: str,
):
    items = [
        ("Исходный файл:", original_filename),
        ("Дата перевода:", datetime.now().strftime("%d.%m.%Y")),
        ("Модель перевода:", model_used),
        ("Метод извлечения:", extraction_method),
    ]
    for label, value in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        lr = p.add_run(f"{label} ")
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.name = "Times New Roman"
        lr.font.color.rgb = RGBColor(80, 80, 80)
        vr = p.add_run(value)
        vr.font.size = Pt(9)
        vr.font.name = "Times New Roman"
        vr.font.color.rgb = RGBColor(80, 80, 80)


def _add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 70)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(180, 180, 180)


def _add_section_heading(doc: Document, label: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(30, 30, 30)


def _add_text_paragraph(doc: Document, text: str):
    """Add one or more paragraphs from a multiline string."""
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def _add_results_table(doc: Document, rows: list[list]):
    """
    Add the test-results table to the document.

    *rows* is a list of lists (first row = header).
    """
    if not rows:
        _add_text_paragraph(doc, "—")
        return

    n_cols = max(len(r) for r in rows)

    # Normalise each row to n_cols
    for row in rows:
        while len(row) < n_cols:
            row.append("")

    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_val in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = str(cell_val) if cell_val else ""

            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)

            # Header row formatting
            if i == 0:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)

    # Spacer after table
    doc.add_paragraph()


def _add_disclaimer(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(
        "Данный документ является переводом оригинального Сертификата анализа.\n"
        "Перевод выполнен с использованием искусственного интеллекта с применением\n"
        "фармацевтического глоссария. Рекомендуется верификация специалистом."
    )
    run.font.size = Pt(8)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(140, 140, 140)
    run.italic = True


def _table_to_text(rows: list[list]) -> str:
    """Convert a table (list of rows) to pipe-delimited text."""
    lines = []
    for row in rows:
        lines.append(" | ".join(str(c) for c in row))
    return "\n".join(lines)


def _template_has_jinja_placeholders(template_bytes: bytes) -> bool:
    """Check whether the template includes expected Jinja placeholders."""
    try:
        doc = Document(io.BytesIO(template_bytes))
    except Exception:
        return False

    text = _extract_document_text(doc)
    placeholder_keys = COA_FIELD_KEYS + [
        "original_filename",
        "translation_date",
        "model_used",
        "extraction_method",
    ]

    for key in placeholder_keys:
        if f"{{{{ {key} }}}}" in text or f"{{{{{key}}}}}" in text:
            return True
    return False


def _rendered_template_has_translated_content(rendered_bytes: bytes, sections: dict) -> bool:
    """
    Verify rendered template contains at least one non-trivial translated
    snippet. Guards against structure-only outputs.
    """
    candidates: list[str] = []
    for key in COA_FIELD_KEYS:
        value = sections.get(key)
        if isinstance(value, str) and len(value.strip()) >= 20:
            candidates.append(value.strip()[:40].lower())
        elif isinstance(value, list) and value:
            first_row = value[1] if len(value) > 1 else value[0]
            if isinstance(first_row, list):
                row_text = " ".join(str(cell) for cell in first_row).strip()
                if row_text:
                    candidates.append(row_text[:40].lower())
        if len(candidates) >= 6:
            break

    if not candidates:
        return True

    try:
        doc = Document(io.BytesIO(rendered_bytes))
    except Exception:
        return False

    rendered_text = _extract_document_text(doc).lower()
    return any(snippet in rendered_text for snippet in candidates)


def _inject_translation_into_template(
    sections: dict,
    original_filename: str,
    extraction_method: str,
    model_used: str,
    template_bytes: bytes,
) -> bytes:
    """
    Fallback for non-Jinja templates:
    1) Try placing content under matching section headings in the template.
    2) Append any remaining translated sections at the end.
    """
    doc = Document(io.BytesIO(template_bytes))

    inserted_keys = _insert_content_under_matching_headings(doc, sections)
    _append_missing_sections(doc, sections, inserted_keys)

    # Ensure metadata is always present in fallback mode.
    _add_horizontal_rule(doc)
    _add_metadata_block(doc, original_filename, extraction_method, model_used)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _insert_content_under_matching_headings(doc: Document, sections: dict) -> set[str]:
    """Insert translated section content under heading-like paragraphs."""
    inserted_keys: set[str] = set()
    paragraphs = list(doc.paragraphs)

    for paragraph in paragraphs:
        key = _match_section_key(paragraph.text)
        if not key or key in inserted_keys:
            continue

        value = sections.get(key)
        if not value:
            continue

        anchor = paragraph
        if isinstance(value, list):
            for row in value:
                if not row:
                    continue
                line = " | ".join(str(cell) for cell in row).strip()
                if not line:
                    continue
                anchor = _insert_paragraph_after(anchor, line)
                _style_body_paragraph(anchor)
        else:
            for line in str(value).splitlines():
                line = line.strip()
                if not line:
                    continue
                anchor = _insert_paragraph_after(anchor, line)
                _style_body_paragraph(anchor)

        inserted_keys.add(key)

    return inserted_keys


def _append_missing_sections(doc: Document, sections: dict, inserted_keys: set[str]):
    """Append non-empty sections not matched by headings."""
    missing = []
    for key, label, _desc, is_table in COA_SECTIONS:
        if key in inserted_keys:
            continue
        value = sections.get(key)
        if not value:
            continue
        missing.append((key, label, is_table, value))

    if not missing:
        return

    _add_horizontal_rule(doc)
    _add_section_heading(doc, "Переведенное содержание")

    for _key, label, is_table, value in missing:
        _add_section_heading(doc, label)
        if is_table and isinstance(value, list):
            _add_results_table(doc, value)
        elif is_table and isinstance(value, str):
            _add_text_paragraph(doc, value)
        else:
            _add_text_paragraph(doc, str(value))


def _match_section_key(text: str) -> str | None:
    """Map paragraph heading text to a known COA section key."""
    heading = _normalise_heading(text)
    if not heading:
        return None

    for key in COA_FIELD_KEYS:
        if heading == key:
            return key

    for key, label in COA_FIELD_LABELS.items():
        if heading == _normalise_heading(label):
            return key

    for key, aliases in SECTION_ALIASES.items():
        for alias in aliases:
            alias_norm = _normalise_heading(alias)
            if not alias_norm:
                continue
            if alias_norm in heading:
                return key
            if SequenceMatcher(a=heading, b=alias_norm).ratio() >= 0.84:
                return key

    for key, label in COA_FIELD_LABELS.items():
        if SequenceMatcher(a=heading, b=_normalise_heading(label)).ratio() >= 0.84:
            return key

    return None


def _insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    """Insert a new paragraph directly after an existing one."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def _style_body_paragraph(paragraph: Paragraph):
    """Apply body text style used in generated sections."""
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(1)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def _extract_document_text(doc: Document) -> str:
    """Read all paragraph and table-cell text from a document."""
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)

    return "\n".join(parts)


def _normalise_heading(text: str) -> str:
    """Normalise headings for robust matching."""
    lowered = text.strip().lower()
    cleaned = re.sub(r"[^a-zа-я0-9 ]+", " ", lowered, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()
